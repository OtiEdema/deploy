import streamlit as st
import PyPDF2
import docx
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain.docstore.document import Document
from transformers import AutoTokenizer, AutoModelForSeq2SeqLM, AutoModelForQuestionAnswering, pipeline

# Function to read and preprocess the document
def read_document(uploaded_file):
    document_text = ""
    try:
        if uploaded_file.type == "text/plain":
            document_text = uploaded_file.read().decode("utf-8")
        elif uploaded_file.type == "application/pdf":
            pdf_reader = PyPDF2.PdfReader(uploaded_file)
            for page in range(len(pdf_reader.pages)):
                document_text += pdf_reader.pages[page].extract_text()
        elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            doc = docx.Document(uploaded_file)
            document_text = "\n".join([para.text for para in doc.paragraphs])
    except Exception as e:
        st.error(f"Error reading document: {e}")
    return document_text

# Streamlit app setup
st.set_page_config(page_title="AI Report Analysis", layout="wide")
st.title("AI-Powered Report Analysis")
st.write("Upload your document to summarize it or ask questions based on its content.")

# File upload component
uploaded_file = st.file_uploader("Choose a file", type=["txt", "pdf", "docx"])

if uploaded_file is not None:
    document_text = read_document(uploaded_file)
    
    # Display the document text in an expandable section
    with st.expander("Show Document Text"):
        st.write(document_text)

    try:
        # Convert document text into a list of Document objects
        documents = [Document(page_content=document_text)]
        
        # Initialize the embedding model
        embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/msmarco-distilbert-base-v3")
        
        # Create a FAISS index for efficient similarity search
        vectorstore = FAISS.from_documents(documents, embeddings)
        
        st.success("FAISS index created successfully")

        # Initialize the summarization model
        tokenizer_summarization = AutoTokenizer.from_pretrained("facebook/bart-large-cnn")
        model_summarization = AutoModelForSeq2SeqLM.from_pretrained("facebook/bart-large-cnn")
        summarization_pipeline = pipeline("summarization", model=model_summarization, tokenizer=tokenizer_summarization)
        
        st.success("Summarization pipeline created successfully")

        # Initialize the QA-specific model
        tokenizer_qa = AutoTokenizer.from_pretrained("deepset/roberta-base-squad2")
        model_qa = AutoModelForQuestionAnswering.from_pretrained("deepset/roberta-base-squad2")
        qa_pipeline = pipeline("question-answering", model=model_qa, tokenizer=tokenizer_qa)
        
        st.success("Question Answering pipeline created successfully")

    except Exception as e:
        st.error(f"Error setting up the analysis components: {e}")
        st.stop()

    # Add a tab selection for Summarization or Question Answering
    tab_selection = st.selectbox("Choose Functionality", ["Summarization", "Question Answering"])

    if tab_selection == "Summarization":
        # Summarization tab
        st.subheader("Summarize Your Report")
        # Summarization percentage slider
        summarization_percentage = st.slider("Select Summarization Percentage", 10, 100, 50)
        
        if st.button("Generate Summary"):
            try:
                # Calculate max_length and min_length based on summarization percentage
                total_length = len(document_text.split())
                max_length = total_length * summarization_percentage // 100
                min_length = max_length // 2
                
                # Retrieve key sections of the report
                retrieved_docs = vectorstore.similarity_search("summary", k=5)
                combined_context = " ".join([doc.page_content for doc in retrieved_docs])

                # Generate a summary using the summarization model
                summary = summarization_pipeline(combined_context, max_length=max_length, min_length=min_length, do_sample=False)

                st.markdown("### Summary")
                st.write(summary[0]['summary_text'])
            except Exception as e:
                st.error(f"Error during summarization: {e}")

    elif tab_selection == "Question Answering":
        # Question Answering tab
        st.subheader("Ask Questions Based on Your Report")

        # Input query from the user
        query = st.text_input("Enter your query:", "")

        if st.button("Get Answer"):
            try:
                # Retrieve relevant documents from the FAISS index
                retrieved_docs = vectorstore.similarity_search(query, k=3)
                combined_context = " ".join([doc.page_content for doc in retrieved_docs])

                # Use the QA pipeline to generate an answer based on the combined context
                qa_input = {
                    'question': query,
                    'context': combined_context
                }
                answer = qa_pipeline(qa_input)
                
                st.markdown("### Answer")
                st.write(answer['answer'])
            except Exception as e:
                st.error(f"Error during retrieval and generation: {e}")
